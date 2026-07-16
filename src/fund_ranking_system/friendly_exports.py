from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


MARKDOWN_REPORTS = [
    ("主分析报告", "fund_analysis_report.md"),
    ("基准与同类对比报告", "benchmark_comparison.md"),
    ("组合构建报告", "portfolio_construction.md"),
    ("组合建议说明书", "portfolio_recommendation.md"),
    ("组合再平衡回测", "portfolio_rebalance_report.md"),
    ("基金级动态权重报告", "adaptive_weight_report.md"),
    ("动态权重 Walk-Forward 验证", "adaptive_backtest_summary.md"),
    ("ML 辅助评分报告", "ml_model_report.md"),
    ("模型效果评估", "ml_evaluation.md"),
    ("原始排名 vs ML 排名对比", "ranking_comparison.md"),
    ("策略回测基准对比", "strategy_benchmark.md"),
    ("数据质量诊断", "data_quality_diagnostics.md"),
    ("基金池准入报告", "fund_universe.md"),
    ("Walk-Forward 验证报告", "backtest_summary.md"),
    ("因子贡献解释", "factor_contributions.md"),
    ("LIME 局部解释", "lime_explanations.md"),
    ("因子相关性诊断", "factor_diagnostics.md"),
    ("权重敏感性报告", "weight_sensitivity.md"),
    ("权重扰动稳健性报告", "weight_robustness.md"),
    ("研究附录", "research_enhancement.md"),
    ("P3 研究报告", "p3_research_enhancement.md"),
]

CSV_REPORTS = [
    ("分析区间", "analysis_window.csv", "reports"),
    ("ML学习权重", "ml_learned_weights.csv", "reports"),
    ("基金级动态权重", "adaptive_factor_weights.csv", "reports"),
    ("动态权重排名", None, "reports"),
    ("基准组合对比", "benchmark_comparison.csv", "reports"),
    ("同类基金对比", None, "reports"),
    ("组合构建摘要", "portfolio_summary.csv", "reports"),
    ("组合持仓权重", None, "reports"),
    ("组合约束配置", "portfolio_constraints.csv", "reports"),
    ("组合建议明细", "portfolio_recommendations.csv", "reports"),
    ("组合风险控制", "portfolio_risk_controls.csv", "reports"),
    ("组合再平衡结果", "portfolio_rebalance_results.csv", "reports"),
    ("组合再平衡窗口", "portfolio_rebalance_periods.csv", "reports"),
    ("动态权重验证结果", "adaptive_walk_forward_results.csv", "reports"),
    ("动态权重验证窗口", "adaptive_walk_forward_periods.csv", "reports"),
    ("ML辅助排名", None, "reports"),
    ("排名对比", None, "reports"),
    ("ML训练样本", "ml_training_samples.csv", "reports"),
    ("模型效果评估", "ml_evaluation.csv", "reports"),
    ("LIME局部解释", "lime_explanations.csv", "reports"),
    ("因子贡献", "factor_contributions.csv", "reports"),
    ("WalkForward结果", "walk_forward_results.csv", "reports"),
    ("WalkForward窗口", "walk_forward_periods.csv", "reports"),
    ("权重敏感性", "weight_sensitivity.csv", "reports"),
    ("权重稳健性", "weight_robustness.csv", "reports"),
    ("基金池准入", "fund_universe.csv", "reports"),
    ("数据质量诊断", "data_quality_diagnostics.csv", "reports"),
    ("策略回测基准", "strategy_benchmark.csv", "reports"),
    ("当前画像排名", None, "reports"),
    ("全部画像排名", "ranking_all_profiles.csv", "processed"),
    ("指标明细", "fund_metrics.csv", "processed"),
]


def save_friendly_exports(
    reports_dir: str | Path,
    processed_dir: str | Path,
    profile: str,
) -> tuple[Path, Path, Path]:
    """Create DOCX, PDF, and XLSX exports for user-friendly downloads."""
    reports_dir = Path(reports_dir)
    processed_dir = Path(processed_dir)
    docx_path = reports_dir / "analysis_reports.docx"
    pdf_path = reports_dir / "analysis_reports.pdf"
    xlsx_path = reports_dir / "analysis_data.xlsx"
    markdown_sections = _load_markdown_sections(reports_dir)
    _write_docx(markdown_sections, docx_path)
    _write_pdf(markdown_sections, pdf_path)
    _write_xlsx(reports_dir, processed_dir, profile, xlsx_path)
    return docx_path, pdf_path, xlsx_path


def _load_markdown_sections(reports_dir: Path) -> list[tuple[str, str]]:
    sections = []
    for title, filename in MARKDOWN_REPORTS:
        path = reports_dir / filename
        if path.exists():
            sections.append((title, path.read_text(encoding="utf-8")))
    if not sections:
        sections.append(("分析报告", "当前没有可用 Markdown 报告。"))
    return sections


def _write_docx(sections: list[tuple[str, str]], path: Path) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    document.add_heading("公募基金风险收益分析综合报告", level=0)
    document.add_paragraph("本文件汇总本次分析生成的主要 Markdown 报告，便于在 Word 中阅读和归档。")

    for index, (title, markdown) in enumerate(sections):
        if index:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        document.add_heading(title, level=1)
        _append_markdown_to_docx(document, markdown)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _append_markdown_to_docx(document: Document, markdown: str) -> None:
    table_buffer: list[list[str]] = []
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if _is_markdown_table_line(line):
            cells = [cell.strip(" `") for cell in line.strip("|").split("|")]
            if not _is_table_separator(cells):
                table_buffer.append(cells)
            continue
        if table_buffer:
            _add_docx_table(document, table_buffer)
            table_buffer = []
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(_strip_markdown(line[2:]), level=1)
        elif line.startswith("## "):
            document.add_heading(_strip_markdown(line[3:]), level=2)
        elif line.startswith("### "):
            document.add_heading(_strip_markdown(line[4:]), level=3)
        elif line.startswith("- "):
            document.add_paragraph(_strip_markdown(line[2:]), style="List Bullet")
        elif line.startswith("```"):
            continue
        else:
            document.add_paragraph(_strip_markdown(line))
    if table_buffer:
        _add_docx_table(document, table_buffer)


def _add_docx_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    header = table.rows[0].cells
    for index in range(column_count):
        header[index].text = rows[0][index] if index < len(rows[0]) else ""
        for paragraph in header[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for index in range(column_count):
            cells[index].text = _strip_markdown(row[index]) if index < len(row) else ""


def _write_pdf(sections: list[tuple[str, str]], path: Path) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    for style in styles.byName.values():
        style.fontName = "STSong-Light"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    story = [
        Paragraph("公募基金风险收益分析综合报告", styles["Title"]),
        Spacer(1, 8),
        Paragraph("本文件汇总本次分析生成的主要报告，便于 PDF 阅读和归档。", styles["BodyText"]),
        Spacer(1, 12),
    ]
    for title, markdown in sections:
        story.append(Paragraph(title, styles["Heading1"]))
        story.extend(_markdown_to_pdf_flowables(markdown, styles))
        story.append(Spacer(1, 10))
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


def _markdown_to_pdf_flowables(markdown: str, styles) -> list[object]:
    flowables: list[object] = []
    table_buffer: list[list[str]] = []
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if _is_markdown_table_line(line):
            cells = [_strip_markdown(cell.strip()) for cell in line.strip("|").split("|")]
            if not _is_table_separator(cells):
                table_buffer.append(cells)
            continue
        if table_buffer:
            flowables.append(_pdf_table(table_buffer))
            flowables.append(Spacer(1, 6))
            table_buffer = []
        if not line or line.startswith("```"):
            continue
        if line.startswith("# "):
            flowables.append(Paragraph(_escape_pdf_text(_strip_markdown(line[2:])), styles["Heading1"]))
        elif line.startswith("## "):
            flowables.append(Paragraph(_escape_pdf_text(_strip_markdown(line[3:])), styles["Heading2"]))
        elif line.startswith("### "):
            flowables.append(Paragraph(_escape_pdf_text(_strip_markdown(line[4:])), styles["Heading3"]))
        elif line.startswith("- "):
            flowables.append(Paragraph("- " + _escape_pdf_text(_strip_markdown(line[2:])), styles["BodyText"]))
        else:
            flowables.append(Paragraph(_escape_pdf_text(_strip_markdown(line)), styles["BodyText"]))
        flowables.append(Spacer(1, 4))
    if table_buffer:
        flowables.append(_pdf_table(table_buffer))
    return flowables


def _pdf_table(rows: list[list[str]]) -> Table:
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows[:16]]
    table = Table(normalized, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF1F5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#17202A")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D9DEE7")),
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def _write_xlsx(reports_dir: Path, processed_dir: Path, profile: str, path: Path) -> None:
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Summary"
    summary.append(["文件", "说明"])
    summary.append(["analysis_reports.docx", "Word 综合报告"])
    summary.append(["analysis_reports.pdf", "PDF 综合报告"])
    summary.append(["analysis_data.xlsx", "Excel 数据汇总"])
    _style_sheet(summary)

    for sheet_name, filename, root_name in CSV_REPORTS:
        actual_filename = filename
        if sheet_name == "ML辅助排名":
            actual_filename = f"ranking_ml_{profile}.csv"
        elif sheet_name == "动态权重排名":
            actual_filename = f"ranking_adaptive_{profile}.csv"
        elif sheet_name == "同类基金对比":
            actual_filename = f"peer_comparison_{profile}.csv"
        elif sheet_name == "组合持仓权重":
            actual_filename = f"portfolio_weights_{profile}.csv"
        elif sheet_name == "排名对比":
            actual_filename = f"ranking_comparison_{profile}.csv"
        elif sheet_name == "当前画像排名":
            actual_filename = f"ranking_{profile}.csv"
        root = reports_dir if root_name == "reports" else processed_dir
        csv_path = root / str(actual_filename)
        if csv_path.exists():
            frame = _safe_read_csv(csv_path)
            if not frame.empty or len(frame.columns) > 0:
                _add_frame_sheet(workbook, sheet_name, frame)

    path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(path)


def _safe_read_csv(path: Path) -> pd.DataFrame:
    try:
        return pd.read_csv(path)
    except pd.errors.EmptyDataError:
        return pd.DataFrame()


def _add_frame_sheet(workbook: Workbook, sheet_name: str, frame: pd.DataFrame) -> None:
    safe_name = _safe_sheet_name(sheet_name, workbook.sheetnames)
    worksheet = workbook.create_sheet(safe_name)
    worksheet.append(list(frame.columns))
    for row in frame.itertuples(index=False):
        worksheet.append(list(row))
    _style_sheet(worksheet)


def _style_sheet(worksheet) -> None:
    header_fill = PatternFill("solid", fgColor="EAF1F5")
    for cell in worksheet[1]:
        cell.font = Font(bold=True, color="17202A")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in worksheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for column in worksheet.columns:
        letter = get_column_letter(column[0].column)
        max_len = max(len(str(cell.value)) if cell.value is not None else 0 for cell in column)
        worksheet.column_dimensions[letter].width = min(max(max_len + 2, 10), 36)
    worksheet.freeze_panes = "A2"
    worksheet.auto_filter.ref = worksheet.dimensions


def _safe_sheet_name(name: str, existing: list[str]) -> str:
    cleaned = re.sub(r"[\[\]:*?/\\]", "", name)[:31] or "Sheet"
    candidate = cleaned
    suffix = 1
    while candidate in existing:
        tail = f"_{suffix}"
        candidate = cleaned[: 31 - len(tail)] + tail
        suffix += 1
    return candidate


def _is_markdown_table_line(line: str) -> bool:
    return line.startswith("|") and line.endswith("|") and "|" in line[1:-1]


def _is_table_separator(cells: list[str]) -> bool:
    return all(set(cell.replace(":", "").replace("-", "").strip()) == set() for cell in cells)


def _strip_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def _escape_pdf_text(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
